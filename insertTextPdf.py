import os
from docx import Document
from openpyxl import load_workbook
from docx.shared import Inches

# Path to the image file
image_path = 'cSign.jpeg'

# Open the Excel workbook and select the sheet
wb = load_workbook(filename='SOA.xlsx')
sheet = wb.active

# Create output folder if it doesn't exist
if not os.path.exists('inserted docs'):
    os.makedirs('inserted docs')

# Iterate over the cells in column E from row 6 to 216
for row in range(6, 217):
    obl_no = sheet['E' + str(row)].value

    # Open the Word document
    doc = Document('EXPDF.docx')

    # Dictionary to store the text to be inserted after each match
    insertions = {
        "OBL NO:": obl_no,
        "CONSIGNEE NAME:": "MEL-BACH ENTERPRISES",
        "CONSIGNEE BANK ACCOUNT NUMBER:": "1023649513",
        "CONSIGNEE BANK NAME:": "UBA BANK PLC",
        "AGENT DETAILS:": "DONCLIMAX BONDED TERMINALS, ONNE",
        "AGENT NAME:": "DONCLIMAX VENTURES",
        "|": "",
        "Tel:": "08136987236    ",
        "E-MAIL:-": "donclimax22@yahoo.com",
        "Name:": "OKAFOR CHIDINMA R.         ",
        "Date:": "24/06/2024"
    }

    # Flags for each key in the insertions dictionary
    inserted = {key: False for key in insertions}

    # Flags to indicate whether to insert the email and telephone number
    insert_tel_no = False

    # Flag to indicate whether the image has been added
    image_added = False

    # Insert text in tables
    tel_no_count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in insertions.items():
                        if key in paragraph.text and not inserted[key]:
                            # Split the paragraph into parts: before key, key itself, and after key
                            parts = paragraph.text.split(key)
                            if len(parts) > 1:
                                # Clear the current paragraph text
                                paragraph.clear()
                                # Add the parts with the key and the value
                                run_before = paragraph.add_run(parts[0])
                                run_key = paragraph.add_run(key)
                                if key == "|" and not image_added:
                                    # Add two spaces and the image
                                    paragraph.clear()
                                    run = paragraph.add_run("  ")
                                    run.add_picture(
                                        image_path, width=Inches(0.5))
                                    image_added = True  # Set the flag to True
                                else:
                                    run_value = paragraph.add_run(f"  {value}")
                                    run_value.bold = True
                                # Add the remaining text after the key
                                if len(parts) > 1:
                                    run_after = paragraph.add_run(parts[1])
                                inserted[key] = True
                    if "Tel NO:" in cell.text:
                        tel_no_count += 1
                        if tel_no_count == 2:
                            insert_tel_no = True
                    elif insert_tel_no:
                        # Create a new run with the phone number and make it bold
                        run = cell.paragraphs[0].add_run(" 08136987236")
                        run.bold = True
                        insert_tel_no = False

    # Save the Word document in the output folder with a unique name based on the OBL number
    doc.save(f'inserted docs/{obl_no}.docx')

print("Word documents created successfully.")
