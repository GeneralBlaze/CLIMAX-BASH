import openpyxl
from docx import Document
from docx.shared import Pt
from collections import defaultdict
import os,re

# Open the Excel file
wb = openpyxl.load_workbook('SOA.xlsx')
sheet = wb.active

# Open the Word document
doc = Document('MSC.docx')

# Create a dictionary to count the number of containers with the same bill of lading number
container_counts = defaultdict(int)

# Iterate over the rows from 6 to 216 to count the containers
for i in range(6, 217):
    bill_of_lading = sheet.cell(row=i, column=5).value
    container_counts[bill_of_lading] += 1

# Create a new folder for the new documents
new_folder = 'REFUND LETTERS'
os.makedirs(new_folder, exist_ok=True)

# Create a dictionary to store the containers for each bill of lading
containers = defaultdict(list)

# Iterate over the rows from 6 to 216 to collect the containers
for i in range(6, 217):
    container = sheet.cell(row=i, column=8).value
    bill_of_lading = sheet.cell(row=i, column=5).value
    containers[bill_of_lading].append(container)

# Iterate over the unique bill of lading numbers to replace the text in the Word document
for bill_of_lading, container_list in containers.items():
    # Skip this iteration if bill_of_lading is None
    if bill_of_lading is None:
        continue

    # Create a string with all the containers separated by commas
    # Ensure all items in container_list are strings before joining
    container_text = ', '.join(str(container) for container in container_list if container is not None)

    # Iterate over the unique bill of lading numbers to replace the text in the Word document
for bill_of_lading, container_list in containers.items():
    # Skip this iteration if bill_of_lading is None
    if bill_of_lading is None:
        continue

    # Create a string with all the containers separated by commas
    # Ensure all items in container_list are strings before joining
    container_text = ', '.join(str(container) for container in container_list if container is not None)

    # Iterate over the rows from 6 to 216 to collect the invoice_no, size, vessel name and voyage number for each container
    for i in range(6, 217):
        current_bill_of_lading = sheet.cell(row=i, column=5).value
        if current_bill_of_lading == bill_of_lading:
            # Get the data from the Excel file
            invoice_no = sheet.cell(row=i, column=16).value
            size = sheet.cell(row=i, column=11).value
            vessel_name = sheet.cell(row=i, column=9).value  # Column I for vessel name
            voyage_no = sheet.cell(row=i, column=10).value  # Column J for voyage number

            # Skip this iteration if invoice_no, size, vessel_name or voyage_no is None
            if invoice_no is None or size is None or vessel_name is None or voyage_no is None:
                continue

            # Calculate the new size text
            size_text = f"({len(container_list)}*{size}FT)"

            # Open the Word document
            doc = Document('MSC.docx')

            # Replace the text in the Word document
            for paragraph in doc.paragraphs:
                if 'CONTAINER' in paragraph.text:
                    # Find the second occurrence of 'CONTAINER' and replace the text after it
                    match = re.search(r'(CONTAINER.*?CONTAINER\s)(.*?)(\s\()', paragraph.text)
                    if match:
                        paragraph.text = paragraph.text.replace(match.group(2), container_text)
                if 'BILL OF LADING:' in paragraph.text:
                    # Replace the text after 'BILL OF LADING:' and before ','
                    match = re.search(r'(BILL OF LADING:\s)(.*?)(,)', paragraph.text)
                    if match:
                        paragraph.text = paragraph.text.replace(match.group(2), bill_of_lading)
                if 'INVOICE NO:' in paragraph.text:
                    # Replace the text after 'INVOICE NO:' and before ' OF VESSEL NAME:'
                    match = re.search(r'(INVOICE NO:\s)(.*?)(\sOF VESSEL NAME:)', paragraph.text)
                    if match:
                        paragraph.text = paragraph.text.replace(match.group(2), str(invoice_no))
                if '(2*20FT)' in paragraph.text:
                    paragraph.text = paragraph.text.replace('(2*20FT)', size_text)
                if 'VESSEL NAME:' in paragraph.text:
                    # Replace the text after 'VESSEL NAME:' and before ','
                    match = re.search(r'(VESSEL NAME:\s)(.*?)(,)', paragraph.text)
                    if match:
                        paragraph.text = paragraph.text.replace(match.group(2), vessel_name)
                if 'VOY -' in paragraph.text:
                    # Replace the text after 'VOY -' and before the next whitespace
                    match = re.search(r'(VOY -\s)(.*?)(\s)', paragraph.text)
                    if match:
                        paragraph.text = paragraph.text.replace(match.group(2), voyage_no)
                    
          # Use regex to find and format the required parts
                for match in re.finditer(r'APPLICATION FOR REFUND OF CONTAINER DEPOSIT MADE ON CONTAINER (?P<containers>[\w, ]+) \((?P<size>\d+\*\d+FT)\) BILL OF LADING: (?P<bill_of_lading>\w+), INVOICE NO: (?P<invoice_no>\w+) OF VESSEL NAME: (?P<vessel_name>[\w\s]+), VOY - (?P<voyage_no>\w+)', paragraph.text):
                    start, end = match.span()
                    found_text = paragraph.text[start:end]
                    for run in paragraph.runs:
                        if run.text in found_text:
                            run.bold = True
                            run.underline = True

            # Save the Word document in the new folder
            doc.save(os.path.join(new_folder, f'MSC REFUND {bill_of_lading} LETTER.docx'))